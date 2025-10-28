from pathlib import Path
import os, argparse, json, time, re, subprocess
import itertools
from typing import List, Dict, Any, Tuple, Optional

from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
from fastapi.responses import PlainTextResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.middleware.cors import CORSMiddleware
import uvicorn
import traceback
from lxml import etree as _ET
import zipfile, io, tempfile, shutil

# --- 3rd party (graceful import) ---
try:
    import requests
except Exception:
    requests = None
try:
    import fasttext
except Exception:
    fasttext = None
try:
    from docx import Document as DocxDocument
except Exception:
    DocxDocument = None
try:
    import openpyxl
    from openpyxl.cell.cell import MergedCell
except Exception:
    openpyxl = None
    class MergedCell: pass
try:
    import ctranslate2
    import sentencepiece as spm
except Exception:
    ctranslate2 = None
    spm = None

APP_VERSION = "0.8.1"

BASE_DIR    = os.path.abspath(os.path.dirname(__file__))
DATA_DIR    = os.path.join(BASE_DIR, "data")
UPLOADS_DIR = os.path.join(DATA_DIR, "uploads")
OUTPUTS_DIR = os.path.join(DATA_DIR, "outputs")
TMP_DIR     = os.path.join(DATA_DIR, "tmp")
LOGS_DIR    = os.path.join(DATA_DIR, "logs")
MODELS_DIR  = os.path.join(BASE_DIR, "models")  # symlink -> /home/udemirci/apps/ai/models

DEFAULT_ENGINE = "nllb"
DEFAULT_NLLB   = "nllb_ct2/600M-int8"
OLLAMA_HOST    = os.environ.get("OLLAMA_HOST", "http://127.0.0.1:11434")

MAX_FILE_MB       = int(os.getenv("MAX_FILE_MB", "100"))
REQUEST_TIMEOUT_S = int(os.getenv("REQUEST_TIMEOUT_S", "900"))
BATCH_SIZE        = int(os.getenv("BATCH_SIZE", "16"))
MAX_INPUT_TOKENS  = int(os.getenv("MAX_INPUT_TOKENS", "900"))
ALLOWED_EXTS      = {".docx", ".xlsx", ".xlsm", ".xltx", ".xltm"}

# ---------- ring log ----------
_RING: List[str] = []
def ringlog(msg: str, maxlen: int = 4000) -> None:
    ts = time.strftime("%Y-%m-%d %H:%M:%S")
    line = f"[{ts}] {msg}"
    _RING.append(line)
    if len(_RING) > maxlen:
        del _RING[:len(_RING)-maxlen]

# ---------- cache ----------
_cache: Dict[str, Any] = {
    "ollama_models": [],
    "ollama_refreshed_at": None,
    "nllb_translators": {},
    "nllb_tokenizers": {},
}

# ---------- languages ----------
SUPPORTED_LANGS = [
    {"code":"auto","name":"Auto-detect"},
    {"code":"en","name":"English"}, {"code":"tr","name":"Turkish"},
    {"code":"fr","name":"French"},  {"code":"de","name":"German"},
    {"code":"es","name":"Spanish"}, {"code":"it","name":"Italian"},
    {"code":"pt","name":"Portuguese"}, {"code":"nl","name":"Dutch"},
    {"code":"ru","name":"Russian"}, {"code":"uk","name":"Ukrainian"},
    {"code":"ar","name":"Arabic"},  {"code":"pl","name":"Polish"},
    {"code":"ro","name":"Romanian"},{"code":"el","name":"Greek"},
    {"code":"cs","name":"Czech"},   {"code":"sk","name":"Slovak"},
    {"code":"bg","name":"Bulgarian"},{"code":"hu","name":"Hungarian"},
    {"code":"sr","name":"Serbian"}, {"code":"hr","name":"Croatian"},
    {"code":"sv","name":"Swedish"}, {"code":"no","name":"Norwegian"},
    {"code":"da","name":"Danish"},  {"code":"fi","name":"Finnish"},
    {"code":"fa","name":"Persian"}, {"code":"zh","name":"Chinese"},
    {"code":"ja","name":"Japanese"},{"code":"ko","name":"Korean"},
]

NLLB_TAGS = {
    "en":"eng_Latn","fr":"fra_Latn","de":"deu_Latn","es":"spa_Latn","it":"ita_Latn",
    "pt":"por_Latn","nl":"nld_Latn","tr":"tur_Latn","ro":"ron_Latn","pl":"pol_Latn",
    "sv":"swe_Latn","no":"nob_Latn","da":"dan_Latn","fi":"fin_Latn","cs":"ces_Latn",
    "sk":"slk_Latn","hr":"hrv_Latn","hu":"hun_Latn",
    "ru":"rus_Cyrl","uk":"ukr_Cyrl","sr":"srp_Cyrl","bg":"bul_Cyrl",
    "ar":"arb_Arab","fa":"pes_Arab",
    "zh":"zho_Hans","ja":"jpn_Jpan","ko":"kor_Hang",
}

# ---------- FastAPI ----------
app = FastAPI(title="Dokuman Ceviri API", version=APP_VERSION)
app.add_middleware(CORSMiddleware, allow_origins=["*"], allow_credentials=True, allow_methods=["*"], allow_headers=["*"])
app.mount("/static", StaticFiles(directory="static"), name="static")

# ---------- model listings ----------
def list_nllb_models() -> List[str]:
    root = os.path.join(MODELS_DIR, "nllb_ct2")
    out = []
    if os.path.isdir(root):
        for name in sorted(os.listdir(root)):
            p = os.path.join(root, name)
            if os.path.isdir(p) and os.path.isfile(os.path.join(p,"model.bin")) and os.path.isfile(os.path.join(p,"sentencepiece.bpe.model")):
                out.append(f"nllb_ct2/{name}")
    return out

def list_argos_models() -> List[str]:
    root = os.path.join(MODELS_DIR, "argos")
    out = []
    if os.path.isdir(root):
        for name in sorted(os.listdir(root)):
            p = os.path.join(root, name)
            if os.path.isdir(p) or name.endswith(".argosmodel"):
                out.append(f"argos/{name}")
    return out

def _ollama_cli_list() -> List[str]:
    try:
        r = subprocess.run(["ollama","list","--json"], check=True, capture_output=True, text=True)
        data = json.loads(r.stdout)
        return sorted([it.get("name") for it in data if it.get("name")])
    except Exception:
        return []

def _ollama_rest_list() -> List[str]:
    if not requests: return []
    try:
        r = requests.get(f"{OLLAMA_HOST}/api/tags", timeout=5)
        r.raise_for_status()
        models = r.json().get("models", [])
        return sorted([m.get("name") for m in models if m.get("name")])
    except Exception:
        return []

def refresh_ollama_models() -> List[str]:
    names = _ollama_cli_list() or _ollama_rest_list()
    _cache["ollama_models"] = names
    _cache["ollama_refreshed_at"] = time.time()
    ringlog(f"Ollama modelleri yenilendi: {len(names)} adet")
    return names

def get_ollama_models_cached() -> List[str]:
    return _cache["ollama_models"] or refresh_ollama_models()

# ---------- LID ----------
def quick_docx_sample_text(path: str, limit_chars: int = 4000) -> str:
    if DocxDocument is None: return ""
    try:
        doc = DocxDocument(path)
        out, s = [], 0
        for p in doc.paragraphs:
            if p.text:
                out.append(p.text); s += len(p.text)
                if s >= limit_chars: break
        return "\n".join(out)
    except Exception:
        return ""

def quick_xlsx_sample_text(path: str, limit_cells: int = 4000) -> str:
    if openpyxl is None: return ""
    try:
        wb = openpyxl.load_workbook(path, data_only=True, read_only=True)
        out, n = [], 0
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for v in row:
                    if isinstance(v, str) and v.strip():
                        out.append(v); n += 1
                        if n >= limit_cells: return "\n".join(out)
        return "\n".join(out)
    except Exception:
        return ""

def _guess_lang_heuristic(text: str) -> str:
    t = (text or "").lower()
    # basit kelime/harf ipuçları
    import re
    # ingilizce: tamamen ASCII + tipik kısa kelimeler
    if not any(ord(c)>127 for c in t) and re.search(r'\b(the|and|is|are|of|to|why|what|where|how)\b', t):
        return 'en'
    if re.search(r'\b(le|la|les|des|de|et|à|est|pas|pour)\b', t) or ('é' in t or 'è' in t or 'ç' in t):
        return 'fr'
    if re.search(r'\b(el|la|los|las|de|del|y|para)\b', t) or ('ñ' in t or '¡' in t or '¿' in t):
        return 'es'
    if re.search(r'\b(ve|ile|bir|için|olarak|de|da|mi)\b', t) or ('ğ' in t or 'ş' in t or 'ı' in t or 'İ' in t):
        return 'tr'
    if re.search(r'\b(der|die|das|und|nicht|für|mit)\b', t) or ('ä' in t or 'ö' in t or 'ß' in t):
        return 'de'
    if re.search(r'\b(il|lo|la|gli|che|per)\b', t) or ('à' in t or 'è' in t or 'ò' in t):
        return 'it'
    if re.search(r'[А-Яа-я]', t):
        return 'ru'
    if re.search(r'[\u0600-\u06FF]', t):
        return 'ar'
    if re.search(r'[\u4E00-\u9FFF]', t):
        return 'zh'
    if re.search(r'[\u3040-\u30FF]', t):
        return 'ja'
    if re.search(r'[\uAC00-\uD7AF]', t):
        return 'ko'
    return 'auto'

def detect_lang_auto(path: str, ext: str) -> str:

    lid_path = os.path.join(MODELS_DIR, "lid.176.ftz")

    if fasttext is None or not os.path.isfile(lid_path):

        ringlog("LID skip: fastText yok; source=auto")

        return "auto"



    sample = quick_docx_sample_text(path) if ext == ".docx" else quick_xlsx_sample_text(path)

    sample = (sample or "").strip()

    if not sample:

        return "auto"



    try:

        ft = fasttext.load_model(lid_path)

        counts = {}

        for line in itertools.islice((l for l in sample.splitlines() if l.strip()), 500):

            pred = ft.predict(line[:2000].replace("\t", " "), k=1)

            labels = pred[0] if isinstance(pred, tuple) and len(pred) >= 1 else list(pred)

            labels = list(labels) if not isinstance(labels, (list, tuple)) else labels

            lbl = (str(labels[0]) if labels else "auto").replace("__label__", "").lower()

            counts[lbl] = counts.get(lbl, 0) + 1

        return max(counts, key=counts.get) if counts else "auto"

    except Exception as e:
        ringlog(f"LID error: {e}")
        h = _guess_lang_heuristic(sample)
        ringlog(f"LID fallback(heur): {h}")
        return h



# ---------- masking ----------
MASK_PATTERNS = [
    (r'\b\d{1,4}([./-])\d{1,2}\1\d{2,4}\b', 'DATE'),
    (r'\b\d{1,2}:\d{2}(:\d{2})?\b', 'TIME'),
    (r'\b\d{1,3}(?:[.,]\d{3})*(?:[.,]\d+)?%?\b', 'NUM'),
    (r'[$€£₺]\s?\d+(?:[.,]\d+)?', 'CUR'),
    (r'\{[^{}]+\}|\{\{[^{}]+\}\}|\[\[[^\[\]]+\]\]', 'PH'),
    (r'https?://\S+|www\.\S+', 'URL'),
    (r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b', 'MAIL'),
]
def mask_protect(text: str) -> Tuple[str, List[Tuple[str,str]]]:
    tokens: List[Tuple[str,str]] = []
    masked = text
    for pat, tag in MASK_PATTERNS:
        def repl(m):
            idx = len(tokens)
            ph = f"__{tag}{idx}__"
            tokens.append((ph, m.group(0)))
            return ph
        masked = re.sub(pat, repl, masked)
    return masked, tokens

def mask_restore(text: str, tokens: List[Tuple[str,str]]) -> str:
    out = text
    for ph, orig in tokens:
        out = out.replace(ph, orig)
    return out

# ---------- NLLB core ----------
def load_nllb(model_name: str):
    if ctranslate2 is None or spm is None:
        raise RuntimeError("ctranslate2/sentencepiece kurulu değil")
    model_dir = os.path.join(MODELS_DIR, model_name)
    spm_path = os.path.join(model_dir, "sentencepiece.bpe.model")
    if model_dir not in _cache["nllb_translators"]:
        ringlog(f"NLLB yükleniyor: {model_dir}")
        tr = ctranslate2.Translator(model_dir, device="cuda", compute_type="int8_float16")
        _cache["nllb_translators"][model_dir] = tr
    else:
        tr = _cache["nllb_translators"][model_dir]
    if model_dir not in _cache["nllb_tokenizers"]:
        sp = spm.SentencePieceProcessor(); sp.load(spm_path)
        _cache["nllb_tokenizers"][model_dir] = sp
    else:
        sp = _cache["nllb_tokenizers"][model_dir]
    return tr, sp

def _clean_nllb(text: str) -> str:
    text = re.sub(r">>\w+<<", " ", text)
    text = re.sub(r"\b\w{3}_(Latn|Cyrl|Arab|Deva|Hans|Hant)\b", " ", text)
    text = re.sub(r"\s{2,}", " ", text).strip()
    # aşırı tekrarları törpüle (aynı kelime 6+)
    text = re.sub(r"(?:\s+(\S+))\1{6,}", lambda m: (" " + m.group(1))*5, " "+text).strip()
    return text

def _split_by_tokens(sp, text: str, max_tokens: int) -> List[str]:
    toks = sp.encode(text, out_type=str)
    if len(toks) <= max_tokens: return [text]
    parts, cur = [], []
    for t in toks:
        cur.append(t)
        if len(cur) >= max_tokens:
            parts.append(sp.decode(cur)); cur=[]
    if cur: parts.append(sp.decode(cur))
    return parts

def _safe_clip(src_text: str, out_text: str, max_mult: float = 3.0) -> str:
    limit = int(max(len(src_text) * max_mult, 1024))
    return out_text[:limit] if len(out_text) > limit else out_text

def nllb_translate_texts(model_name: str, src_lang: str, tgt_lang: str, texts: List[str]) -> List[str]:
    tr, sp = load_nllb(model_name)
    tgt_tag = NLLB_TAGS.get((tgt_lang or "en").lower(), "eng_Latn")
    src_tag = NLLB_TAGS.get((src_lang or "auto").lower(), None) if (src_lang and src_lang != "auto") else None
    tgt_tok = sp.encode(tgt_tag, out_type=str)

    def encode_input(s: str) -> List[str]:
        s = s or ""
        if src_tag:
            s = f"{src_tag} {s}"
        return sp.encode(s, out_type=str)

    tokenized = [encode_input(t) for t in texts]
    results: List[str] = []
    i = 0
    while i < len(texts):
        batch_tokens: List[List[str]] = []
        idxs: List[int] = []
        while i < len(texts) and len(batch_tokens) < BATCH_SIZE:
            toks = tokenized[i]
            if len(toks) > MAX_INPUT_TOKENS:
                parts = _split_by_tokens(sp, texts[i], MAX_INPUT_TOKENS - (20 if src_tag else 0))
                pieces = []
                for ptxt in parts:
                    ptoks = encode_input(ptxt)
                    out = tr.translate_batch(
                        [ptoks],
                        beam_size=1,
                        length_penalty=1.2,
                        repetition_penalty=1.3,
                        no_repeat_ngram_size=8,
                        max_decoding_length=256,
                        target_prefix=[tgt_tok],
                        return_alternatives=False,
                        disable_unk=True,
                    )
                    tok_out = out[0].hypotheses[0]
                    piece = sp.decode(tok_out)
                    pieces.append(_clean_nllb(piece))
                joined = " ".join(pieces).strip()
                results.append(_safe_clip(texts[i], joined))
                i += 1
                continue
            batch_tokens.append(toks); idxs.append(i); i += 1

        if batch_tokens:
            prefixes = [tgt_tok]*len(batch_tokens)
            ringlog(f"NLLB batch translate {idxs[0]}..{idxs[-1]} / {len(texts)} -> {tgt_tag}" + (f" (src={src_tag})" if src_tag else ""))
            out = tr.translate_batch(
                batch_tokens,
                beam_size=1,
                length_penalty=1.2,
                repetition_penalty=1.3,
                no_repeat_ngram_size=8,
                max_decoding_length=256,
                target_prefix=prefixes,
                return_alternatives=False,
                disable_unk=True,
            )
            for j,o in enumerate(out):
                decoded = sp.decode(o.hypotheses[0])
                cleaned = _clean_nllb(decoded)
                results.append(_safe_clip(texts[idxs[j]], cleaned))
    return results

# ---------- Argos (optional) ----------
def argos_translate_batch(texts: List[str], src: str, tgt: str) -> List[str]:
    try:
        import argostranslate.translate as argos_tr
    except Exception as e:
        raise RuntimeError(f"Argos import error: {e}")
    ringlog(f"ARGOS batch: src={src} tgt={tgt} items={len(texts)}")
    outs = []
    for i, t in enumerate(texts):
        try:
            this_src = src
            if (src or "auto") == "auto":
                # parça bazlı tespit
                this_src = _guess_lang_heuristic(t)
            if not this_src or this_src == "auto" or this_src == tgt:
                # kaynak bilinmiyorsa veya hedefle aynıysa çeviri anlamsız—pas geç
                out = t or ""
            else:
                out = argos_tr.translate(t or "", this_src, tgt)
                # çok karışık belgelerde yanlış kaynakla aynı çıktıyı veriyorsa bir kez heuristik dene
                if (not out or out.strip() == (t or "").strip()) and src != "auto":
                    h = _guess_lang_heuristic(t)
                    if h and h != "auto" and h != tgt and h != this_src:
                        try:
                            out2 = argos_tr.translate(t or "", h, tgt)
                            # bir şey değiştiyse onu al
                            if out2 and out2.strip() != (t or "").strip():
                                ringlog(f"ARGOS retry[{i}]: {this_src}->{tgt} -> {h}->{tgt}")
                                out = out2
                        except Exception:
                            pass
            outs.append(out)
        except Exception:
            outs.append(t or "")
    return outs

# ---------- Post-Edit (optional) ----------
def ollama_post_edit(model: str, text: str) -> str:
    if not requests: return text
    try:
        prompt = ("You are a professional translator. Improve fluency and terminology consistency. "
                  "Do not add or remove information. Keep placeholders and numbers intact.\n\nText:\n"+text)
        r = requests.post(f"{OLLAMA_HOST}/api/generate", json={"model":model,"prompt":prompt,"stream":False}, timeout=120)
        r.raise_for_status()
        data = r.json()
        return (data.get("response") or "").strip() or text
    except Exception:
        return text

# ---------- Sentence splitting ----------
_SENT_SPLIT_RX = re.compile(r"(\n+|[.!?;:•·]+[\s]+)")
def _split_sentences_keep_seps(text: str) -> List[str]:
    if not text: return []
    parts = _SENT_SPLIT_RX.split(text)
    return parts if parts else [text]

# ---------- Pipeline ----------
def translate_texts_pipeline_native(texts: List[str], engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]) -> List[str]:
    ringlog(f"PIPE native: engine={engine} model={model_name or '-'} src={src} tgt={tgt} n={len(texts)} post_edit={post_edit or '-'}")
    # 1) mask
    masked_list, banks_list = [], []
    for t in texts:
        m, bank = mask_protect(t or "")
        masked_list.append(m); banks_list.append(bank)

    # 2) sentence-wise flatten
    all_chunks: List[str] = []
    owners: List[Tuple[int,int]] = []  # (text_idx, part_idx)
    chunk_is_sep: List[bool] = []
    for idx, m in enumerate(masked_list):
        parts = _split_sentences_keep_seps(m)
        for j, part in enumerate(parts):
            is_sep = True if (j % 2 == 1) else False  # 0:text,1:sep,2:text,3:sep...
            owners.append((idx, j))
            chunk_is_sep.append(is_sep)
            if is_sep or not (part or "").strip():
                all_chunks.append(part)
            else:
                all_chunks.append(part)

    # 3) translate only text chunks in batches
    def do_engine(texts_batch: List[str]) -> List[str]:
        if engine == "nllb":
            return nllb_translate_texts(model_name, src, tgt, texts_batch)
        elif engine == "argos":
            return argos_translate_batch(texts_batch, src, tgt)
        else:
            raise HTTPException(status_code=400, detail=f"Unknown engine: {engine}")

    # prepare indices
    to_tr_idx = [i for i, is_sep in enumerate(chunk_is_sep) if not is_sep and (all_chunks[i] or "").strip()]
    outs_map: Dict[int,str] = {}
    p = 0
    while p < len(to_tr_idx):
        q = min(p + BATCH_SIZE, len(to_tr_idx))
        batch_ids = to_tr_idx[p:q]
        batch_txt = [all_chunks[i] for i in batch_ids]
        trs = do_engine(batch_txt)
        if post_edit:
            ringlog(f"POST-EDIT: model={post_edit} items={len(trs)}")
            trs = [ollama_post_edit(post_edit, x) for x in trs]
        for k, ii in enumerate(batch_ids):
            outs_map[ii] = trs[k]
        p = q

    # 4) reassemble and unmask per original text
    agg: Dict[int, List[str]] = {}
    for i,(owner_text_idx, part_idx) in enumerate(owners):
        s = outs_map[i] if i in outs_map else all_chunks[i]
        agg.setdefault(owner_text_idx, []).append(s)

    out_all: List[str] = []
    for idx in range(len(masked_list)):
        joined = "".join(agg.get(idx, []))
        restored = mask_restore(joined, banks_list[idx])
        out_all.append(restored)
    return out_all

# ---------- DOCX ----------
def _iter_all_paragraphs(doc):
    for p in doc.paragraphs: yield p
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs: yield p
    for sec in doc.sections:
        for part in (sec.header, sec.footer):
            if not part: continue
            for p in part.paragraphs: yield p
            for tbl in part.tables:
                for row in tbl.rows:
                    for cell in row.cells:
                        for p in cell.paragraphs: yield p

def translate_docx(in_path: str, out_path: str, engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]) -> None:
    if DocxDocument is None:
        raise RuntimeError("python-docx yüklü değil")
    ringlog("DOCX açılıyor…")
    doc = DocxDocument(in_path)
    processed_par_elems = set()
    paras = list(_iter_all_paragraphs(doc))
    ringlog(f"DOCX paragraphs total={len(paras)}")

    idxs, payload = [], []
    for i,p in enumerate(paras):
        txt = p.text or ""
        if txt.strip():
            idxs.append(i); payload.append(txt)

    if not payload:
        ringlog("DOCX: çevrilecek metin yok; kopyalanıyor")
        doc.save(out_path); return

    outs = translate_texts_pipeline(payload, engine, model_name, src, tgt, post_edit)

    for k, pi in enumerate(idxs):
        p = paras[pi]
        new_txt = outs[k]
        if len(p.runs)==0:
            p.add_run(new_txt)
        else:
            p.runs[0].text = new_txt
            for r in p.runs[1:]:
                r.text = ""
        try:
            processed_par_elems.add(p._element)
        except Exception:
            pass
        if k % 50 == 0:
            ringlog(f"DOCX para {k}/{len(idxs)}")
    # TextBox (DrawingML) içerikleri de çevir
    try:
        translate_docx_textboxes2(doc, engine, model_name, src, tgt, post_edit)
    except Exception as _e:
        ringlog(f"TextBox translate error: {_e}")
    # DrawingML a:txBody (Text Box) içeriklerini de çevir
    try:
        translate_docx_draw_textbodies(doc, engine, model_name, src, tgt, post_edit)
    except Exception as _e:
        ringlog(f"DML TextBody translate error: {_e}")
    # GENEL drawing/text kapsayıcılarını da çevir
    try:
        translate_docx_any_drawing_text(doc, engine, model_name, src, tgt, post_edit)
    except Exception as _e:
        ringlog(f"DRAW/TEXT translate error: {_e}")
    # SDT (content control) paragraflarını da çevir
    try:
        translate_docx_sdt_paragraphs(doc, engine, model_name, src, tgt, post_edit)
    except Exception as _e:
        ringlog(f"SDT translate error: {_e}")
    # Genel fallback: kalan tüm w:t paragraflarını (işlenmemiş) çevir
    try:
        translate_docx_all_wt_fallback(doc, engine, model_name, src, tgt, post_edit, processed_par_elems)
    except Exception as _e:
        ringlog(f"FALLBACK w:t error: {_e}")

    ringlog("DOCX kaydediliyor…")
    doc.save(out_path)
# ---------- XLSX ----------
def build_merged_maps(ws):
    masters, children = set(), set()
    for r in ws.merged_cells.ranges:
        mr, mc = r.min_row, r.min_col
        masters.add((mr, mc))
        for (rr, cc) in r.cells:
            if (rr, cc) != (mr, mc):
                children.add((rr, cc))
    return masters, children

def translate_xlsx(in_path: str, out_path: str, engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]) -> None:
    if openpyxl is None:
        raise RuntimeError("openpyxl yüklü değil")
    ringlog("XLSX açılıyor…")
    wb = openpyxl.load_workbook(in_path)
    for ws in wb.worksheets:
        masters, children = build_merged_maps(ws)
        batch_texts: List[str] = []
        batch_cells: List[Any] = []
        count = 0
        for row in ws.iter_rows():
            for cell in row:
                if (cell.row, cell.column) in children:  # merged child
                    continue
                if isinstance(cell, MergedCell):
                    continue
                if cell.data_type == 'f' or (isinstance(cell.value, str) and cell.value.startswith("=")):
                    continue
                if getattr(cell, "is_date", False) or isinstance(cell.value, (int, float)):
                    continue
                if not isinstance(cell.value, str) or not cell.value.strip():
                    continue
                batch_texts.append(cell.value); batch_cells.append(cell); count += 1
                if len(batch_texts) >= BATCH_SIZE:
                    ringlog(f"XLSX sheet={ws.title} batch={len(batch_texts)}")
                    outs = translate_texts_pipeline(batch_texts, engine, model_name, src, tgt, post_edit)
                    for c, t in zip(batch_cells, outs): c.value = t
                    batch_texts, batch_cells = [], []
        if batch_texts:
            ringlog(f"XLSX sheet={ws.title} tail-batch={len(batch_texts)}")
            outs = translate_texts_pipeline(batch_texts, engine, model_name, src, tgt, post_edit)
            for c, t in zip(batch_cells, outs): c.value = t
        ringlog(f"XLSX sheet={ws.title} translated_cells={count}")
    ringlog("XLSX kaydediliyor…"); wb.save(out_path)

# ---------- routes ----------
@app.get("/healthz", response_class=PlainTextResponse)
def healthz(): return "ok"

@app.get("/", response_class=HTMLResponse)
def root():
    try:
        with open("static/index.html","r",encoding="utf-8") as f: return f.read()
    except Exception:
        return "<html><body><h3>Doküman Çeviri API</h3></body></html>"

@app.get("/languages")
def languages(): return JSONResponse(SUPPORTED_LANGS)

@app.get("/ollama/refresh")
def ollama_refresh():
    names = refresh_ollama_models()
    return {"count": len(names), "models": names, "refreshed_at": _cache["ollama_refreshed_at"]}

@app.get("/models")
def models():
    return {
        "nllb": list_nllb_models(),
        "argos": list_argos_models(),
        "ollama": get_ollama_models_cached(),
        "default_engine": DEFAULT_ENGINE,
        "default_model": DEFAULT_NLLB,
    }

@app.get("/logs")
def logs(limit: int = 200): return {"lines": _RING[-limit:]}

@app.post("/logs/clear")
def logs_clear():
    try:
        _RING.clear()
        return {"ok": True}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"logs clear error: {e}")

@app.post("/translate")
async def translate(
    file: UploadFile = File(...),
    source_lang: str = Form("auto"),
    target_lang: str = Form("en"),
    engine: str = Form("nllb"),
    model: str = Form(DEFAULT_NLLB),
    post_edit_model: Optional[str] = Form(None),
):
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in ALLOWED_EXTS:
        raise HTTPException(status_code=400, detail=f"Unsupported file extension: {ext}")

    os.makedirs(UPLOADS_DIR, exist_ok=True)
    os.makedirs(OUTPUTS_DIR, exist_ok=True)

    safe_name = os.path.basename(file.filename or "input")
    up_path = os.path.join(UPLOADS_DIR, safe_name)

    ringlog(f"Yükleme başladı: {safe_name}")
    with open(up_path, "wb") as f:
        while True:
            chunk = file.file.read(1024*1024)
            if not chunk: break
            f.write(chunk)
    ringlog(f"Yükleme bitti: {up_path}")
    
    # Boyut limiti (MAX_FILE_MB) kontrolü
    try:
        sz = os.path.getsize(up_path)
        if sz > (MAX_FILE_MB * 1024 * 1024):
            ringlog(f"Dosya sınırı aşıldı: {sz} > {MAX_FILE_MB}MB")
            os.remove(up_path)
            raise HTTPException(status_code=413, detail="file too large")
    except HTTPException:
        raise
    except Exception as _e:
        ringlog(f"size check error: {_e}")

    src = (source_lang or "auto").lower().strip()
    if src == "auto":
        src = detect_lang_auto(up_path, ext)
        ringlog(f"Auto-detect sonucu: {src}")
    if engine == "argos":
        model = ""
    ringlog(f"ENGINE: engine={engine} model={(model or '-')} post_edit={(post_edit_model or '-')} src={src} tgt={target_lang}")

    out_name = f"{os.path.splitext(safe_name)[0]} [{(src or 'auto').upper() if src!='auto' else 'AUTO'}-{target_lang.upper()}]{ext}"
    out_path = os.path.join(OUTPUTS_DIR, out_name)

    t0 = time.time()
    try:
        if ext == ".docx":
            translate_docx(up_path, out_path, engine, model, src, target_lang, post_edit_model)
        else:
            translate_xlsx(up_path, out_path, engine, model, src, target_lang, post_edit_model)
    except Exception as e:
        try:
            import traceback
            ringlog("Çeviri hatası (traceback başlayor)")
            for line in traceback.format_exc().splitlines():
                ringlog(line)
        except Exception:
            pass
        ringlog(f"Çeviri hatası (özet): {e}")
        raise HTTPException(status_code=500, detail=f"translate failed: {e}")

    dur = time.time() - t0
    ringlog(f"Çıktı hazır: {out_path} (süre: {dur:.1f}s)")
    return {
        "status":"done","engine":engine,"model":model,"post_edit_model":post_edit_model,
        "source_lang":src,"target_lang":target_lang,
        "input_file":safe_name,"output_file":out_name,"output_path":out_path,
        "duration_s":round(dur,1),
    }

# ---------- boot ----------
def _ollama_warmup():
    try: refresh_ollama_models()
    except Exception: pass

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--serve", action="store_true")
    ap.add_argument("--port", type=int, default=int(os.getenv("PORT","8085")))
    args = ap.parse_args()
    if args.serve:
        _ollama_warmup()
        uvicorn.run("ceviri:app", host="0.0.0.0", port=args.port, workers=1, access_log=False)
    else:
        print("Use --serve to start API")

if __name__ == "__main__":
    main()


# === Ollama LLM çeviri (engine="ollama") ===
def _chunk_text(txt: str, max_chars: int = 1200) -> list:
    txt = (txt or "").strip()
    if not txt:
        return []
    parts, cur, cur_len = [], [], 0
    for line in txt.splitlines():
        line = line.strip()
        if not line:
            if cur:
                parts.append(" ".join(cur)); cur, cur_len = [], 0
            continue
        if cur_len + len(line) + 1 > max_chars:
            if cur:
                parts.append(" ".join(cur))
            cur, cur_len = [line], len(line)
        else:
            cur.append(line); cur_len += len(line) + 1
    if cur:
        parts.append(" ".join(cur))
    return parts

def _ollama_build_prompt(text: str, source_lang: str, target_lang: str) -> str:
    src = source_lang or "auto"
    return (
        f"Translate the following text from {src} to {target_lang}. "
        f"Return only the translated text, without explanations or formatting metadata:\n\n{text}"
    )

def translate_with_ollama_api(texts: list, model: str, source_lang: str, target_lang: str) -> list:
    ringlog(f"Ollama API: model={model} src={source_lang} tgt={target_lang} texts={len(texts)}")
    if not requests:
        raise RuntimeError("requests modülü yok (Ollama için gerekli)")
    url = f"{OLLAMA_HOST}/api/generate"
    outs = []
    for t in texts:
        chunks = _chunk_text(t, max_chars=1200)
        translated = []
        for ch in chunks:
            payload = {
                "model": model,
                "prompt": _ollama_build_prompt(ch, source_lang, target_lang),
                "stream": False,
                "options": {"temperature": 0.2, "num_ctx": 2048},
            }
            r = requests.post(url, json=payload, timeout=600)
            r.raise_for_status()
            data = r.json()
            translated.append((data.get("response") or "").strip())
        outs.append("\n".join(translated).strip())
    return outs


def translate_texts_pipeline(texts: List[str], engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]) -> List[str]:
    ringlog(f"PIPE: engine={engine} model={model_name or '-'} src={src} tgt={tgt} n={len(texts)} post_edit={post_edit or '-'}")
    """
    Wrapper:
      - engine == "ollama" => LLM ile çeviri (maskeli), main.py yaklaşımı
      - diğerleri          => mevcut yerel akış (NLLB/Argos)
    """
    if engine == "ollama":
        # Maskeleri uygula
        masked_list, banks_list = [], []
        for t in texts:
            m, bank = mask_protect(t or "")
            masked_list.append(m); banks_list.append(bank)

        # LLM çevirisi
        llm_outs = translate_with_ollama_api(masked_list, model_name, src, tgt)

        # (opsiyonel) Post-edit
        if post_edit:
            ringlog(f"POST-EDIT: model={post_edit} items={len(llm_outs)}")
            llm_outs = [ollama_post_edit(post_edit, x) for x in llm_outs]

        # Maskeleri geri koy
        out_all = []
        for i, out in enumerate(llm_outs):
            out_all.append(mask_restore(out, banks_list[i]))
        return out_all

    # Default: mevcut native akış
    return translate_texts_pipeline_native(texts, engine, model_name, src, tgt, post_edit)

# === DOCX TextBox (DrawingML) çevirisi ===
try:
    from docx.oxml.ns import nsmap as _DOCX_NSMAP
except Exception:
    _DOCX_NSMAP = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}

def _iter_doc_parts_for_txbx(document):
    """Belge part'larını (ana gövde + header/footer part'ları) güvenle üretir."""
    # Ana belge
    try:
        yield document.part
    except Exception:
        pass
    # Header/Footer part'ları
    try:
        for sec in getattr(document, "sections", []):
            try:
                hdr = getattr(sec, "header", None)
                if hdr and getattr(hdr, "part", None):
                    yield hdr.part
            except Exception:
                pass
            try:
                ftr = getattr(sec, "footer", None)
                if ftr and getattr(ftr, "part", None):
                    yield ftr.part
            except Exception:
                pass
    except Exception:
        pass

def _collect_dml_txbody_groups(root, ns):
    """
    DrawingML metin kutularını (a:txBody) topla ve metin nodelarıyla birlikte döndür.
    return: [(t_nodes_list, joined_text), ...]
    """
    groups = []
    try:
        for tx in root.findall(".//a:txBody", ns):
            t_nodes = tx.xpath(".//a:p//a:r//a:t", namespaces=ns)
            if not t_nodes:
                continue
            txt = "".join([(t.text or "") for t in t_nodes]).strip()
            if not txt:
                continue
            groups.append((t_nodes, txt))
    except Exception:
        return []
    return groups
    # Kimi DOCX'lerde iliştirilmiş ek header/footer part'ları olabilir:
    try:
        relparts = getattr(document.part, "related_parts", {})
        for rp in getattr(relparts, "values", lambda: [])():
            if hasattr(rp, "element"):
                yield rp
    except Exception:
        pass

def _collect_txbx_paragraph_elements(document):
    """Tüm part'larda w:txbxContent//w:p elementlerini döndürür (lxml elemanları)."""
    paras = []
    for part in _iter_doc_parts_for_txbx(document):
        try:
            elems = part.element.xpath(".//w:txbxContent//w:p", namespaces=_DOCX_NSMAP)
            if elems:
                paras.extend(elems)
        except Exception:
            continue
    return paras

def translate_docx_textboxes(document, engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]):
    """TextBox içi paragrafları çevirir (w:t düğümleri birleştirilip tek metin)."""
    try:
        paras = _collect_txbx_paragraph_elements(document)
    except Exception as e:
        ringlog(f"TextBox scan error: {e}")
        return

    if not paras:
        ringlog("TextBox: bulunamadı")
        return

    texts = []
    groups = []  # her paragraf için w:t düğüm listesi
    for p in paras:
        try:
            t_nodes = p.xpath(".//w:t", namespaces=_DOCX_NSMAP)
            if not t_nodes:
                continue
            raw = "".join([t.text or "" for t in t_nodes]).strip()
            if not raw:
                continue
            texts.append(raw)
            groups.append(t_nodes)
        except Exception:
            continue

    if not texts:
        ringlog("TextBox: çevrilecek metin yok")
        return

    ringlog(f"TextBox paragraphs to_translate={len(texts)}")
    outs = translate_texts_pipeline(texts, engine, model_name, src, tgt, post_edit)

    # Geri yaz
    for t_nodes, new_txt in zip(groups, outs):
        try:
            # ilk w:t'ye çeviriyi yaz, diğerlerini boşalt
            t_nodes[0].text = new_txt
            for tn in t_nodes[1:]:
                tn.text = ""
        except Exception:
            continue

    ringlog("TextBox: çeviri tamamlandı")

# === Genişletilmiş DOCX TextBox (DrawingML + VML) çevirisi ===
def _get_nsmap():
    # python-docx nsmap eksikse kendimiz tamamlayalım
    ns = {
        "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "v":   "urn:schemas-microsoft-com:vml",
        "mc":  "http://schemas.openxmlformats.org/markup-compatibility/2006",
    }
    try:
        from docx.oxml.ns import nsmap as _DOCX_NSMAP
        ns.update({k:v for k,v in _DOCX_NSMAP.items() if k and v})
    except Exception:
        pass
    return ns

def _iter_doc_parts_for_txbx2(document):
    yield document.part
    try:
        for sect in document.sections:
            for part in (sect.header, sect.footer):
                if part is not None and hasattr(part, "part"):
                    yield part.part
    except Exception:
        pass
    # related parts (ör. ek header/footer veya embedded)
    try:
        relparts = getattr(document.part, "related_parts", {})
        for rp in getattr(relparts, "values", lambda: [])():
            if hasattr(rp, "element"):
                yield rp
    except Exception:
        pass

def _collect_txbx_paragraph_elements2(document):
    """Tüm part'larda (DrawingML + VML) text box paragraflarını topla."""
    ns = _get_nsmap()
    xpaths = [
        ".//w:txbxContent//w:p",   # genel
        ".//wps:txbx//w:p",        # Word 2010 shapes
        ".//v:textbox//w:p",       # legacy VML textbox
    ]
    paras = []
    for part in _iter_doc_parts_for_txbx2(document):
        for xp in xpaths:
            try:
                found = part.element.xpath(xp, namespaces=ns)
                if found:
                    paras.extend(found)
            except Exception:
                continue
    return paras

def translate_docx_textboxes2(document, engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]):
    """TextBox içi paragrafları çevir (birleştir, çevir, ilk w:t'ye yaz)."""
    try:
        paras = _collect_txbx_paragraph_elements2(document)
    except Exception as e:
        ringlog(f"TextBox scan error (2): {e}")
        return

    if not paras:
        ringlog("TextBox(2): bulunamadı")
        return

    texts = []
    groups = []   # her paragraf için w:t düğümleri
    ns = _get_nsmap()
    for p in paras:
        try:
            t_nodes = p.xpath(".//w:t", namespaces=ns)
            if not t_nodes:
                continue
            raw = "".join([(t.text or "") for t in t_nodes]).strip()
            if not raw:
                continue
            texts.append(raw)
            groups.append(t_nodes)
        except Exception:
            continue

    if not texts:
        ringlog("TextBox(2): çevrilecek metin yok")
        return

    ringlog(f"TextBox(2) paragraphs to_translate={len(texts)}")
    outs = translate_texts_pipeline(texts, engine, model_name, src, tgt, post_edit)

    # Geri yaz
    for t_nodes, new_txt in zip(groups, outs):
        try:
            t_nodes[0].text = new_txt
            for tn in t_nodes[1:]:
                tn.text = ""
        except Exception:
            continue

    ringlog("TextBox(2): çeviri tamamlandı")

# === DOCX DrawingML Text Body (a:txBody) çevirisi ===
def _nsmap_full():
    ns = {
        "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "v":   "urn:schemas-microsoft-com:vml",
        "mc":  "http://schemas.openxmlformats.org/markup-compatibility/2006",
    }
    try:
        from docx.oxml.ns import nsmap as _DOCX_NSMAP  # type: ignore
        ns.update({k:v for k,v in _DOCX_NSMAP.items() if k and v})
    except Exception:
        pass
    return ns

def _iter_all_parts(document):
    # Main document
    yield document.part
    # Section header/footer
    try:
        for sect in document.sections:
            for part in (sect.header, sect.footer):
                if part is not None and hasattr(part, "part"):
                    yield part.part
    except Exception:
        pass
    # Related parts (embedded/extra)
    try:
        relparts = getattr(document.part, "related_parts", {})
        for rp in getattr(relparts, "values", lambda: [])():
            if hasattr(rp, "element"):
                yield rp
    except Exception:
        pass

def _collect_dml_textbodies(document):
    """Tüm part'larda a:txBody düğümlerini toplar."""
    ns = _nsmap_full()
    bodies = []
    for part in _iter_all_parts(document):
        try:
            # a:txBody genelde wp:inline/wp:anchor altında gelir
            found = part.element.xpath(".//a:txBody", namespaces=ns)
            if found:
                bodies.extend(found)
        except Exception:
            continue
    return bodies

def translate_docx_draw_textbodies(document, engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]):
    """
    Her a:txBody içinde yer alan a:p/a:r/a:t düğümlerini toplayıp tek metin haline getirir,
    çevirir ve ilk a:t'ye yazar, diğerlerini boşaltır.
    """
    ns = _nsmap_full()
    bodies = _collect_dml_textbodies(document)
    if not bodies:
        ringlog("DML TextBody: bulunamadı")
        return

    texts = []
    groups = []  # her body için (t_nodes listesi)
    for body in bodies:
        try:
            t_nodes = body.xpath(".//a:p//a:r//a:t", namespaces=ns)
            if not t_nodes:
                continue
            raw = "".join([(t.text or "") for t in t_nodes]).strip()
            if not raw:
                continue
            texts.append(raw)
            groups.append(t_nodes)
        except Exception:
            continue

    if not texts:
        ringlog("DML TextBody: çevrilecek metin yok")
        return

    ringlog(f"DML TextBody to_translate={len(texts)}")
    outs = translate_texts_pipeline(texts, engine, model_name, src, tgt, post_edit)

    # Geri yaz
    for t_nodes, new_txt in zip(groups, outs):
        try:
            t_nodes[0].text = new_txt
            for tn in t_nodes[1:]:
                tn.text = ""
        except Exception:
            continue

    ringlog("DML TextBody: çeviri tamamlandı")

# === GENEL: DOCX içindeki tüm drawing/textbox kapsayıcılarını yakala ve çevir ===
def _docx_nsmap_all():
    ns = {
        "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
        "v":   "urn:schemas-microsoft-com:vml",
        "mc":  "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
    }
    try:
        from docx.oxml.ns import nsmap as _DOCX_NSMAP  # type: ignore
        ns.update({k:v for k,v in _DOCX_NSMAP.items() if k and v})
    except Exception:
        pass
    return ns

def _iter_all_parts_for_draw(document):
    # Main
    yield document.part
    # Section header/footer
    try:
        for sect in document.sections:
            for part in (sect.header, sect.footer):
                if part is not None and hasattr(part, "part"):
                    yield part.part
    except Exception:
        pass
    # Related parts
    try:
        relparts = getattr(document.part, "related_parts", {})
        for rp in getattr(relparts, "values", lambda: [])():
            if hasattr(rp, "element"):
                yield rp
    except Exception:
        pass

def _collect_generic_text_containers(document):
    """
    Kapsayıcı düğümler:
      - w:drawing
      - w:pict
      - mc:AlternateContent
      - w:txbxContent
      - wps:txbx
      - v:textbox
      - a:txBody (DrawingML)
    Her kapsayıcı içinde w:t düğümlerini toplayacağız.
    """
    ns = _docx_nsmap_all()
    container_xpaths = [
        ".//w:drawing",
        ".//w:pict",
        ".//mc:AlternateContent",
        ".//w:txbxContent",
        ".//wps:txbx",
        ".//v:textbox",
        ".//a:txBody",
    ]
    containers = []
    for part in _iter_all_parts_for_draw(document):
        el = part.element
        for xp in container_xpaths:
            try:
                found = el.xpath(xp, namespaces=ns)
                if found:
                    containers.extend(found)
            except Exception:
                continue
    # Aynı düğümü tekrarlamayalım
    uniq = []
    seen = set()
    for c in containers:
        key = id(c)
        if key in seen: 
            continue
        seen.add(key)
        uniq.append(c)
    return uniq

def translate_docx_any_drawing_text(document, engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]):
    ns = _docx_nsmap_all()
    containers = _collect_generic_text_containers(document)
    if not containers:
        ringlog("DRAW/TEXT: kapsayıcı bulunamadı")
        return

    groups = []  # her kapsayıcı için w:t düğüm listesi
    texts  = []
    for c in containers:
        try:
            # Bu kapsayıcı içindeki tüm w:t düğümlerini sırayla al
            t_nodes = c.xpath(".//w:t", namespaces=ns)
            if not t_nodes:
                continue
            raw = "".join([(t.text or "") for t in t_nodes]).strip()
            if not raw:
                continue
            texts.append(raw)
            groups.append(t_nodes)
        except Exception:
            continue

    if not texts:
        ringlog("DRAW/TEXT: çevrilecek metin yok")
        return

    ringlog(f"DRAW/TEXT groups to_translate={len(texts)}")
    outs = translate_texts_pipeline(texts, engine, model_name, src, tgt, post_edit)

    # Geri yaz: her grup tek satır
    for t_nodes, new_txt in zip(groups, outs):
        try:
            t_nodes[0].text = new_txt
            for tn in t_nodes[1:]:
                tn.text = ""
        except Exception:
            continue

    ringlog("DRAW/TEXT: çeviri tamamlandı")

# === DOCX Content Control (SDT) içindeki paragrafları çevir ===
def _nsmap_all():
    ns = {
        "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
        "v":   "urn:schemas-microsoft-com:vml",
        "mc":  "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
    }
    try:
        from docx.oxml.ns import nsmap as _DOCX_NSMAP  # type: ignore
        ns.update({k:v for k,v in _DOCX_NSMAP.items() if k and v})
    except Exception:
        pass
    return ns

def _iter_all_parts_for_sdt(document):
    yield document.part
    try:
        for sect in document.sections:
            for part in (sect.header, sect.footer):
                if part is not None and hasattr(part, "part"):
                    yield part.part
    except Exception:
        pass
    try:
        relparts = getattr(document.part, "related_parts", {})
        for rp in getattr(relparts, "values", lambda: [])():
            if hasattr(rp, "element"):
                yield rp
    except Exception:
        pass

def _collect_sdt_paragraphs(document):
    ns = _nsmap_all()
    paras = []
    for part in _iter_all_parts_for_sdt(document):
        try:
            # SDT content: w:sdt//w:sdtContent//w:p
            found = part.element.xpath(".//w:sdt//w:sdtContent//w:p", namespaces=ns)
            if found:
                paras.extend(found)
        except Exception:
            continue
    return paras

def translate_docx_sdt_paragraphs(document, engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]):
    ns = _nsmap_all()
    paras = _collect_sdt_paragraphs(document)
    if not paras:
        ringlog("SDT: bulunamadı")
        return

    texts, groups = [], []   # her paragraf için w:t düğüm grubu
    for p in paras:
        try:
            t_nodes = p.xpath(".//w:r//w:t", namespaces=ns) or p.xpath(".//w:t", namespaces=ns)
            if not t_nodes:
                continue
            raw = "".join([(t.text or "") for t in t_nodes]).strip()
            if not raw:
                continue
            texts.append(raw)
            groups.append(t_nodes)
        except Exception:
            continue

    if not texts:
        ringlog("SDT: çevrilecek metin yok")
        return

    ringlog(f"SDT paragraphs to_translate={len(texts)}")
    outs = translate_texts_pipeline(texts, engine, model_name, src, tgt, post_edit)

    for t_nodes, new_txt in zip(groups, outs):
        try:
            t_nodes[0].text = new_txt
            for tn in t_nodes[1:]:
                tn.text = ""
        except Exception:
            continue

    ringlog("SDT: çeviri tamamlandı")

# === GENEL FALLBACK: Belgedeki kalan tüm w:t düğümlerini (işlenmemiş paragraflar) çevir ===
def _ns_all_docx():
    ns = {
        "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
        "v":   "urn:schemas-microsoft-com:vml",
        "mc":  "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
    }
    try:
        from docx.oxml.ns import nsmap as _DOCX_NSMAP  # type: ignore
        ns.update({k:v for k,v in _DOCX_NSMAP.items() if k and v})
    except Exception:
        pass
    return ns

def _iter_all_parts_generic(document):
    # Main
    yield document.part
    # Section header/footer
    try:
        for sect in document.sections:
            for part in (sect.header, sect.footer):
                if part is not None and hasattr(part, "part"):
                    yield part.part
    except Exception:
        pass
    # Related parts
    try:
        relparts = getattr(document.part, "related_parts", {})
        for rp in getattr(relparts, "values", lambda: [])():
            if hasattr(rp, "element"):
                yield rp
    except Exception:
        pass

def _nearest_w_p(elem, ns):
    """Verilen XML elem için en yakın üst w:p'yi bul."""
    cur = elem
    while cur is not None:
        if cur.tag.endswith("}p") or cur.tag == "{%s}p" % ns["w"]:
            return cur
        cur = cur.getparent() if hasattr(cur, "getparent") else None
    return None

def translate_docx_all_wt_fallback(document, engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str], processed_par_elems: set):
    """
    İşlenmemiş tüm paragrafların (w:p) içindeki w:t'leri grupla, çevir, geri yaz.
    processed_par_elems: daha önce normal akışta çevirdiğimiz w:p elementleri (dokunma).
    """
    ns = _ns_all_docx()
    groups = []  # [(p_elem, [t_nodes])]
    seen_p = set()

    for part in _iter_all_parts_generic(document):
        try:
            t_nodes = part.element.xpath(".//w:t", namespaces=ns)
        except Exception:
            t_nodes = []
        for tn in t_nodes:
            p = _nearest_w_p(tn, ns)
            if p is None: 
                continue
            if p in processed_par_elems:
                # Bu paragraf zaten normal DOCX akışında çevrildi
                continue
            if id(p) in seen_p:
                continue
            # Bu paragrafın tüm t'lerini topla
            try:
                p_t_nodes = p.xpath(".//w:r//w:t", namespaces=ns) or p.xpath(".//w:t", namespaces=ns)
            except Exception:
                p_t_nodes = []
            if not p_t_nodes:
                continue
            groups.append((p, p_t_nodes))
            seen_p.add(id(p))

    if not groups:
        ringlog("FALLBACK w:t: işlenecek paragraf yok")
        return

    texts = []
    for _, tnodes in groups:
        raw = "".join([(t.text or "") for t in tnodes]).strip()
        if not raw:
            texts.append("")  # boş kalabilir
        else:
            texts.append(raw)

    # Boşlar varsa yine de çeviri çağrısı düzgün çalışsın
    ringlog(f"FALLBACK w:t: groups={len(groups)}")
    outs = translate_texts_pipeline(texts, engine, model_name, src, tgt, post_edit)

    # Geri yaz
    changed = 0
    for (_, tnodes), new_txt in zip(groups, outs):
        if not tnodes:
            continue
        try:
            tnodes[0].text = new_txt or ""
            for tn in tnodes[1:]:
                tn.text = ""
            changed += 1
        except Exception:
            continue
    ringlog(f"FALLBACK w:t: çeviri tamamlandı (paragraf={changed})")

# === ZIP FALLBACK: DOCX içindeki tüm XML'lerde w:t düğümlerini tarayıp çevir ===
def _ns_zip_all():
    return {
        "w":   "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a":   "http://schemas.openxmlformats.org/drawingml/2006/main",
        "wp":  "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
        "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
        "wps": "http://schemas.microsoft.com/office/word/2010/wordprocessingShape",
        "wpg": "http://schemas.microsoft.com/office/word/2010/wordprocessingGroup",
        "v":   "urn:schemas-microsoft-com:vml",
        "mc":  "http://schemas.openxmlformats.org/markup-compatibility/2006",
        "dgm": "http://schemas.openxmlformats.org/drawingml/2006/diagram",
    }

_FRENCH_HINT_RX = re.compile(r"\b(le|la|les|des|de|du|d’|l’|au|aux|un|une|et|ou|dans|avec|pour|sur|par|en)\b", re.IGNORECASE)
_NONASCII_RX    = re.compile(r"[^\x00-\x7F]")

def _looks_french_or_nonascii(s: str, src_lang: str) -> bool:
    if not s or not s.strip():
        return False
    if _NONASCII_RX.search(s):
        return True
    if (src_lang or "auto").lower() in ("fr","auto"):
        return bool(_FRENCH_HINT_RX.search(s))
    return False

def _zip_read_all_word_xml(zf):
    # word/ altındaki tüm .xml dosyalarını getir
    xml_entries = []
    for name in zf.namelist():
        if name.startswith("word/") and name.endswith(".xml"):
            xml_entries.append(name)
    return xml_entries

def translate_docx_zip_fallback(out_path: str, engine: str, model_name: str, src: str, tgt: str, post_edit: Optional[str]):
    ns = _ns_zip_all()
    try:
        with zipfile.ZipFile(out_path, "r") as zf:
            names = _zip_read_all_word_xml(zf)
            mem = {n: zf.read(n) for n in zf.namelist()}

        # 1) DrawingML a:txBody grupları
        dml_groups = []
        parsed = {}
        for name in names:
            try:
                root = _ET.fromstring(mem[name])
            except Exception:
                continue
            parsed[name] = root
            groups = _collect_dml_txbody_groups(root, ns)
            for t_nodes, text in groups:
                if _looks_french_or_nonascii(text, src):
                    dml_groups.append((name, t_nodes, text))

        if dml_groups:
            ringlog(f"ZIP-DML: txBody groups to_translate={len(dml_groups)}")
            texts = [g[2] for g in dml_groups]
            outs  = translate_texts_pipeline(texts, engine, model_name, src, tgt, post_edit)
            for (name, t_nodes, _), new_txt in zip(dml_groups, outs):
                joined = (new_txt or "")
                if t_nodes:
                    t_nodes[0].text = joined
                    for tn in t_nodes[1:]:
                        tn.text = ""
            for name in {n for (n, *_ ) in dml_groups}:
                xml_bytes = _ET.tostring(parsed[name], encoding="utf-8", xml_declaration=True, standalone=False)
                mem[name] = xml_bytes
        else:
            ringlog("ZIP-DML: txBody group bulunamadı")

        # 2) w:t fallback
        wt_nodes, wt_names, wt_texts = [], [], []
        for name in names:
            if name in parsed:
                root = parsed[name]
            else:
                try:
                    root = _ET.fromstring(mem[name])
                except Exception:
                    continue
                parsed[name] = root
            t_nodes = root.xpath(".//w:t", namespaces=ns)
            for tn in t_nodes:
                txt = (tn.text or "").strip()
                if _looks_french_or_nonascii(txt, src):
                    wt_nodes.append(tn); wt_names.append(name); wt_texts.append(txt)

        if wt_texts:
            ringlog(f"ZIP-w:t: to_translate={len(wt_texts)}")
            outs = translate_texts_pipeline(wt_texts, engine, model_name, src, tgt, post_edit)
            for tn, new_txt in zip(wt_nodes, outs):
                tn.text = new_txt
            for name in set(wt_names):
                xml_bytes = _ET.tostring(parsed[name], encoding="utf-8", xml_declaration=True, standalone=False)
                mem[name] = xml_bytes
        else:
            ringlog("ZIP-w:t: çevrilecek düğüm yok")

        tmp_out = out_path + ".tmpzip"
        with zipfile.ZipFile(tmp_out, "w", zipfile.ZIP_DEFLATED) as zfw:
            for n, data in mem.items():
                zfw.writestr(n, data)
        shutil.move(tmp_out, out_path)
        ringlog("ZIP-FALLBACK: yazıldı (DML + w:t)")
    except Exception as e:
        ringlog(f"ZIP-FALLBACK error: {e}")
    ns = _ns_zip_all()
    try:
        with zipfile.ZipFile(out_path, "r") as zf:
            names = _zip_read_all_word_xml(zf)
            # Bellekte kopyasını hazırlayacağız (Zip dosyası in-place editlemez)
            mem = {}
            for n in zf.namelist():
                mem[n] = zf.read(n)

        total_candidates = 0
        total_changed = 0
        # Tüm w:t metinlerini toplayıp batch çeviri yapacağız; ancak aşırı büyük dosyalarda
        # parça parça işleyelim:
        batch_nodes = []  # (name, tree, node)
        batch_texts = []
        parsed_cache = {}  # name -> (_ET.ElementTree, root)
        for name in names:
            try:
                tree = _ET.fromstring(mem[name])
            except Exception:
                continue
            parsed_cache[name] = (tree, tree)
            # XPath: tüm w:t düğümleri
            t_nodes = tree.xpath(".//w:t", namespaces=ns)
            for tn in t_nodes:
                txt = (tn.text or "").strip()
                if _looks_french_or_nonascii(txt, src):
                    batch_nodes.append((name, tn))
                    batch_texts.append(txt)
            total_candidates += len(t_nodes)

        if not batch_texts:
            ringlog("ZIP-FALLBACK: çevrilecek w:t bulunamadı")
            return

        ringlog(f"ZIP-FALLBACK: candidates={len(batch_texts)}/{total_candidates} (word/*.xml)")
        # Batch çeviri (mevcut pipeline)
        outs = translate_texts_pipeline(batch_texts, engine, model_name, src, tgt, post_edit)

        # Çeviriyi yerine yaz
        for (name, tn), new_txt in zip(batch_nodes, outs):
            try:
                tn.text = new_txt
                total_changed += 1
            except Exception:
                continue

        # Yeni zip'i yaz
        tmp_out = out_path + ".tmpzip"
        with zipfile.ZipFile(tmp_out, "w", zipfile.ZIP_DEFLATED) as zfw:
            for n, data in mem.items():
                if n in parsed_cache:
                    root = parsed_cache[n][0]
                    xml_bytes = _ET.tostring(root, encoding="utf-8", xml_declaration=True, standalone=False)
                    zfw.writestr(n, xml_bytes)
                else:
                    zfw.writestr(n, data)
        # Yerine koy
        shutil.move(tmp_out, out_path)
        ringlog(f"ZIP-FALLBACK: yazıldı (changed={total_changed})")
    except Exception as e:
        ringlog(f"ZIP-FALLBACK error: {e}")


@app.get("/download")
def download(file: str):
    """
    Güvenli indirme:
    - Sadece OUTPUTS_DIR altından dosya verir
    - file paramı yalnızca dosya adı (out.output_file) olmalı
    """
    try:
        base = Path(OUTPUTS_DIR).resolve()
        # Sadece dosya adını al (üst dizin kaçışı engelle)
        fname = os.path.basename(file)
        if not fname:
            raise HTTPException(status_code=400, detail="empty filename")
        p = (base / fname).resolve()

        # Base dışına kaçmayı engelle
        if not str(p).startswith(str(base) + os.sep):
            ringlog(f"DOWNLOAD blocked: resolved outside base -> {p}")
            raise HTTPException(status_code=400, detail="invalid path")

        if not p.exists() or not p.is_file():
            ringlog(f"DOWNLOAD not found: {p}")
            raise HTTPException(status_code=404, detail="not found")

        # İçerik türü
        name = p.name.lower()
        if name.endswith(".docx"):
            media = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        elif name.endswith(".xlsx"):
            media = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
        else:
            media = "application/octet-stream"

        ringlog(f"DOWNLOAD ok: {p}")
        return FileResponse(str(p), media_type=media, filename=p.name)
    except HTTPException:
        raise
    except Exception as e:
        ringlog(f"DOWNLOAD error: {e}")
        ringlog(traceback.format_exc())
        raise HTTPException(status_code=500, detail=f"download error: {e}")
